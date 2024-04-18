from docx import Document
import pandas as pd

list_ = []  # 初始化一个空列表，用来装后面的数据字典


# 处理khz
path = "khz.docx"
docx = Document(path)
for table in docx.tables:  # 循环所有的表格
    row_counter = 0  # 初始化行计数器
    for row in table.rows:  # 循环表格中的所有行
        row_counter += 1  # 每次循环时，增加行计数器
        if row_counter <= 2:  # 如果是前两行，则跳过
            continue
        cells = row.cells  # 获取当前行的所有单元格
        # 假设每个单元格中的内容都是通过换行符分隔的
        # 并且第一列包含了起始频率和结束频率，用"—"分隔
        # 其余列包含了业务信息
        if cells:
            item = cells[0].text
            dict_ = {
                '起始频率(KHz)': str(item.split("\n")[0]).split("—")[0],
                '结束频率(KHz)': str(item.split("\n")[0]).split("—")[1],
                '业务': ','.join(item.split("\n")[1:])
            }
            list_.append(dict_)  # 将字典添加到列表中


# Mhz,需要成1000
# 处理khz
path = "mhz.docx"
docx = Document(path)
for table in docx.tables:  # 循环所有的表格
    row_counter = 0  # 初始化行计数器
    for row in table.rows:  # 循环表格中的所有行
        row_counter += 1  # 每次循环时，增加行计数器
        if row_counter <= 2:  # 如果是前两行，则跳过
            continue
        cells = row.cells  # 获取当前行的所有单元格
        # 假设每个单元格中的内容都是通过换行符分隔的
        # 并且第一列包含了起始频率和结束频率，用"—"分隔
        # 其余列包含了业务信息
        if cells:
            item = cells[0].text
            temp=str(item.split("\n")[0]).replace(' ', '')
            dict_ = {
                '起始频率(KHz)': float(temp.split("—")[0])*1000,
                '结束频率(KHz)': float(temp.split("—")[1])*1000,
                '业务': ','.join(item.split("\n")[1:])
            }
            list_.append(dict_)  # 将字典添加到列表中


# GHZ 需要乘以1000 *1000
path = "Ghz.docx"
docx = Document(path)
for table in docx.tables:  # 循环所有的表格
    row_counter = 0  # 初始化行计数器
    for row in table.rows:  # 循环表格中的所有行
        row_counter += 1  # 每次循环时，增加行计数器
        if row_counter <= 2:  # 如果是前两行，则跳过
            continue
        cells = row.cells  # 获取当前行的所有单元格
        # 假设每个单元格中的内容都是通过换行符分隔的
        # 并且第一列包含了起始频率和结束频率，用"—"分隔
        # 其余列包含了业务信息
        if cells:
            item = cells[0].text
            temp=str(item.split("\n")[0]).replace(' ', '')
            dict_ = {
                '起始频率(KHz)': float(temp.split("—")[0])*1000*1000,
                '结束频率(KHz)': float(temp.split("—")[1])*1000*1000,
                '业务': ','.join(item.split("\n")[1:])
            }
            list_.append(dict_)  # 将字典添加到列表中

# 将列表转换为DataFrame
df = pd.DataFrame(list_)

# 将DataFrame保存为Excel文件
excel_path = "alldata.xlsx"
df.to_excel(excel_path, index=False)

print(f"数据已保存至 {excel_path}")